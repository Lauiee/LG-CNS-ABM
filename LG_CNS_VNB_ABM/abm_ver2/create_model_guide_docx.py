#!/usr/bin/env python3
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUT_PATH = BASE_DIR / "PRISM_ABM_시뮬레이션_모델_사용_가이드.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(95, 105, 120)
GRAY_FILL = "F2F4F7"
BLUE_GRAY_FILL = "E8EEF5"
LIGHT_FILL = "F7F9FC"
WHITE = "FFFFFF"
TEXT = RGBColor(25, 31, 40)
RISK_RED = RGBColor(155, 28, 28)
GOOD_GREEN = RGBColor(32, 111, 78)
FONT = "Malgun Gothic"
MONO = "Consolas"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, color=None, bold=None, italic=None, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size=None, color=None, bold=None, font=FONT):
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:ascii"), font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="CBD3DF", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_header_repeat(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def format_cell_text(cell, size=9.2, bold=False, color=TEXT, font=FONT, align=None):
    for para in cell.paragraphs:
        if align is not None:
            para.alignment = align
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.05
        for run in para.runs:
            set_run_font(run, size=size, color=color, bold=bold, font=font)


def add_para(doc, text="", style=None, size=11, color=TEXT, bold=False, italic=False,
             align=None, before=0, after=6, line=1.1):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = add_para(doc, text, style=style, size=10.5, after=4, line=1.12)
    return p


def add_number(doc, text):
    return add_para(doc, text, style="List Number", size=10.5, after=4, line=1.12)


def add_code(doc, text):
    p = add_para(doc, before=0, after=3, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=9.2, color=RGBColor(42, 49, 60), font=MONO)
    p.paragraph_format.left_indent = Inches(0.18)
    return p


def add_callout(doc, title, body, fill=LIGHT_FILL, title_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color="D7DEE9", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=title_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.8, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows, widths, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_header_repeat(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(header)
        set_cell_shading(cell, GRAY_FILL)
        format_cell_text(cell, size=8.8, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
            format_cell_text(cells[i], size=font_size, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        set_style_font(style, size=10.5, color=TEXT)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.12

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "PRISM ABM 시뮬레이션 모델 가이드"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "LG CNS VNB Developer Productivity ABM | 보고 및 공유용"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in fp.runs:
        set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc):
    add_para(doc, "PRISM ABM", size=12, color=BLUE, bold=True, after=6)
    title = add_para(
        doc,
        "시뮬레이션 모델\n사용 및 수정 가이드",
        size=26,
        color=NAVY,
        bold=True,
        after=8,
        line=1.05,
    )
    title.paragraph_format.keep_with_next = True
    add_para(
        doc,
        "Developer Productivity Agent-Based Model 설계 구조, 실행 방법, 실험 파라미터, 수정 지점 정리",
        size=12.5,
        color=MUTED,
        after=18,
    )

    metadata = [
        ("대상 코드", "LG_CNS_VNB_ABM/abm_ver2"),
        ("문서 목적", "모델 이해, 공유, 실행, 실험 설계, 코드 수정 기준 제공"),
        ("주요 사용자", "연구자, PM/PL, 운영자, 개발자, 분석 담당자"),
        ("작성 기준일", date.today().strftime("%Y.%m.%d")),
    ]
    table = doc.add_table(rows=0, cols=2)
    for label, value in metadata:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    set_table_geometry(table, [1800, 7560])
    set_table_borders(table, color="D7DEE9")
    for row in table.rows:
        set_cell_shading(row.cells[0], BLUE_GRAY_FILL)
        format_cell_text(row.cells[0], size=9.5, bold=True, color=NAVY)
        format_cell_text(row.cells[1], size=9.5, color=TEXT)

    add_callout(
        doc,
        "문서 사용법",
        "처음 읽는 사용자는 1~4장을 먼저 보면 모델 구조와 실행 흐름을 이해할 수 있습니다. "
        "실험 담당자는 5~7장의 파라미터 표와 시나리오 표를, 코드 수정 담당자는 8장을 기준으로 수정 범위를 잡으면 됩니다.",
        fill="F4F6F9",
    )
    doc.add_page_break()


def add_manual_toc(doc):
    doc.add_heading("문서 구성", level=1)
    toc_items = [
        "1. 모델 개요와 사용 목적",
        "2. 디렉토리와 파일 구조",
        "3. 설계 구조: 에이전트, 태스크, 환경",
        "4. 실행 흐름과 관계식",
        "5. 초기 세팅과 기본 분포",
        "6. 사용 방법: 웹 UI, 단일 실행, 반복 실험",
        "7. 실험 파라미터 카탈로그",
        "8. 수정 방법과 확장 체크리스트",
        "9. 산출물 해석과 공유 시 유의사항",
        "부록. 빠른 참조 표",
    ]
    for item in toc_items:
        add_bullet(doc, item)


def section_overview(doc):
    doc.add_heading("1. 모델 개요와 사용 목적", level=1)
    add_callout(
        doc,
        "핵심 요약",
        "본 시뮬레이션은 개발 조직을 Developer Agent와 PL Agent의 상호작용으로 표현하고, "
        "회의 부하, 리뷰 엄격도, 코드베이스 안정성, 요구사항 명확도, 백로그 크기, 멘토링, PM 개입 수준 같은 정책 변수가 "
        "PRISM 생산성 지표에 어떤 영향을 주는지 사전 실험하는 정책 실험실입니다.",
    )
    add_para(
        doc,
        "모델은 실제 조직의 모든 맥락을 완전히 재현하기보다, 개발자 상태 변화와 업무 흐름 사이의 방향성 있는 관계를 "
        "반복 실험 가능한 형태로 단순화합니다. 따라서 단일 실행 결과보다 조건별 반복 실행의 평균과 민감도 비교가 중요합니다.",
    )
    add_table(
        doc,
        ["구분", "내용"],
        [
            ("모델 유형", "Mesa 기반 Agent-Based Model"),
            ("시뮬레이션 단위", "1 step = 1일, 1 sprint = 10 step, 기본 6 sprint = 60 step"),
            ("핵심 에이전트", "DeveloperAgent, PLAgent"),
            ("업무 객체", "Task: coding, reviewing, testing, deploying, incident"),
            ("주요 출력", "PRISM 6개 지표 + 내부 진단 지표"),
            ("권장 활용", "정책 변수 조정에 따른 상대 변화, 민감도, trade-off 탐색"),
        ],
        [1900, 7460],
    )

    doc.add_heading("PRISM 출력 지표", level=2)
    add_table(
        doc,
        ["지표", "계산 방식", "해석"],
        [
            ("PRs per Engineer", "활성 개발자 기준 PR 생성 수", "개인당 처리량 지표"),
            ("Lead Time (steps)", "완료/배포까지 평균 소요 step", "낮을수록 빠른 전달"),
            ("Deployment Frequency", "스프린트당 배포 횟수", "높을수록 빈번한 전달"),
            ("Change Failure Rate (%)", "실패/인시던트 발생 비율", "낮을수록 안정적"),
            ("Recovery Time (steps)", "incident 생성부터 완료까지 평균 step", "낮을수록 복구력 높음"),
            ("% Time on New Capabilities", "신규 기능 관련 step 비율", "기능 개발 집중도"),
        ],
        [2100, 3100, 4160],
    )


def section_structure(doc):
    doc.add_heading("2. 디렉토리와 파일 구조", level=1)
    add_para(
        doc,
        "실제 분석 대상은 저장소의 `LG_CNS_VNB_ABM/abm_ver2` 폴더입니다. "
        "모델 본체, 실행 서버, 실험 러너, 보고서 산출물이 같은 폴더 아래에 모여 있습니다.",
    )
    add_table(
        doc,
        ["파일/폴더", "역할", "주로 수정하는 경우"],
        [
            ("model.py", "전체 모델 상태, 전역 파라미터, 스프린트/step 실행 순서, PM 개입, 압박 로직", "환경 변수, step 순서, PM/백로그/회의 관계식을 바꿀 때"),
            ("agents.py", "DeveloperAgent와 PLAgent의 속성, 상태 전이, 업무 수행, 멘토링, 배정 로직", "개발자 행동, PL 행동, 생산성 함수, 멘토링 규칙을 바꿀 때"),
            ("tasks.py", "Task 타입, 난이도 C1~C5, 도메인, incident 우선순위", "업무 분포, 난이도, 결함률, 도메인을 바꿀 때"),
            ("sampling.py", "초기 속성 분포와 distribution_overrides 처리", "초기값을 실측 기반 분포로 보정할 때"),
            ("simulate.py", "단일 시뮬레이션 실행, snapshot, PRISM/내부 지표 반환", "출력 지표를 추가하거나 API 입력을 확장할 때"),
            ("experiment_runner.py", "A~E 시나리오 조건, 반복 실행, CSV 요약", "실험 설계를 추가/수정할 때"),
            ("run_final_experiments.py", "A~E 최종 100회 반복 실행 및 결과 보존", "전체 최종 실험을 재생산할 때"),
            ("server.py / dashboard.html", "웹 UI와 /simulate API", "사용자용 슬라이더나 화면 표시를 바꿀 때"),
            ("outputs/final", "최종 시나리오 결과 CSV, 보고서용 표/그림", "보고서 작성 또는 결과 검토 시"),
        ],
        [2100, 4300, 2960],
        font_size=8.2,
    )

    add_callout(
        doc,
        "주의",
        "웹 UI에 노출된 파라미터와 실험 스크립트에서 조작 가능한 파라미터가 완전히 같지는 않습니다. "
        "팀 구성, 멘토링, PM 프로파일, 백로그 크기 등은 주로 API 또는 experiment_runner.py를 통해 조작합니다.",
        fill="FFF8E8",
        title_color=RGBColor(122, 90, 0),
    )


def section_design(doc):
    doc.add_heading("3. 설계 구조: 에이전트, 태스크, 환경", level=1)
    doc.add_heading("3.1 전체 구조", level=2)
    add_table(
        doc,
        ["구성요소", "상태/속성", "행동", "연결되는 지표"],
        [
            ("DeveloperAgent", "skill, role, energy, motivation, knowledge, domain_knowledge, flow_streak", "코딩, 리뷰, 학습, 도움 요청, 멘토링 수신/제공, 번아웃", "PR, Lead Time, CFR, Recovery, 내부 멘토링 지표"),
            ("PLAgent", "leadership_style, team_awareness, energy, motivation, team_stress", "스탠드업, 번아웃 코칭, incident 배정, 업무 배정, sprint retro", "coaching_count, backlog pressure 완화, attrition"),
            ("Task", "type, complexity, domain, progress, priority, defect_prob", "개발자에게 배정되고 progress가 1.0에 도달하면 리뷰/배포/incident로 전이", "Lead Time, Deployment, CFR, New Capability"),
            ("Environment", "meeting_load, review_strictness, codebase_stability, requirement_clarity, backlog size", "전역 압박, 결함 발생, 요구사항 조율, 지식 감쇠", "평균 에너지, 실패율, 처리량, 잔여 백로그"),
        ],
        [1800, 3150, 2900, 1510],
        font_size=8.1,
    )

    doc.add_heading("3.2 Developer Agent", level=2)
    add_table(
        doc,
        ["속성군", "주요 변수", "의미"],
        [
            ("역량/역할", "skill_level, role", "개발자의 기본 생산성 계수와 junior/middle/senior 역할"),
            ("상태", "energy, motivation, knowledge, flow_streak", "매 step 변화하는 생산성 기반 상태"),
            ("품질/중단", "quality_tendency, interrupt_sensitivity", "기술부채 누적과 회의/인터럽트 취약성"),
            ("도메인", "domain_knowledge", "frontend/backend/data/infra/legacy/testing별 숙련도"),
            ("협업", "help_seeking_tendency, mentoring_tendency, mentoring_load", "도움 요청 가능성과 도움 제공 부담"),
        ],
        [1700, 3200, 4460],
    )
    add_para(doc, "생산성 함수는 다음 곱셈 구조입니다.", after=2)
    add_code(doc, "productivity = skill_coeff * (energy / 100) * (motivation / 100) * (knowledge / 100)")
    add_para(
        doc,
        "skill_coeff는 skill_level별 생산성 계수이며, energy는 작업 가능 자원, motivation은 수행 의지, knowledge는 코드베이스/업무 이해도를 의미합니다. "
        "세 요소 중 하나가 낮아지면 생산성이 함께 낮아지도록 설계되어 있습니다.",
    )

    doc.add_heading("3.3 Developer 상태", level=2)
    add_table(
        doc,
        ["상태", "진입/유지 조건", "효과"],
        [
            ("Idle", "현재 배정된 task 없음", "20% 확률로 Learning 전이"),
            ("Coding", "task 배정 후 기본 작업 상태", "progress 증가, energy 소모, tech debt 누적"),
            ("FlowCoding", "flow_streak >= 3, energy > 60", "몰입 상태. energy 소모가 상대적으로 낮지만 중단 시 streak 초기화"),
            ("Reviewing", "PL이 pending review task를 배정", "review_strictness에 따라 비용/속도/결함 유출이 변함"),
            ("Learning", "Idle에서 자율 학습", "energy 소모, knowledge 증가"),
            ("Interrupted", "incident 또는 interrupt 수신", "energy 소모, flow_streak 초기화 후 복귀"),
            ("Burnout", "energy <= burnout_threshold", "5 step 지속 시 attrition 처리"),
        ],
        [1400, 3500, 4460],
        font_size=8.5,
    )

    doc.add_heading("3.4 PL Agent", level=2)
    add_bullet(doc, "매 step Daily Standup을 수행하여 PL과 팀원의 energy를 소폭 감소시키고 flow_streak를 초기화합니다.")
    add_bullet(doc, "energy가 burnout_threshold * 1.5 아래로 내려간 개발자를 team_awareness 확률로 감지하여 coaching합니다.")
    add_bullet(doc, "incident task는 우선순위에 따라 개발자에게 배정하고, Critical/High incident는 더 강한 interruption을 유발합니다.")
    add_bullet(doc, "일반 task는 유휴 개발자에게 배정합니다. PM 개입이 켜지면 domain knowledge, skill fit, energy, mentoring load를 반영해 best-fit 배정을 시도합니다.")
    add_bullet(doc, "Sprint Retro에서는 평균 energy가 낮으면 backlog size를 줄이고, motivation이 낮으면 review strictness를 완화합니다.")

    doc.add_heading("3.5 Task 설계", level=2)
    add_table(
        doc,
        ["Task 타입", "역할", "주요 영향"],
        [
            ("coding", "신규 기능/작업 개발", "PR 생성, Lead Time 시작"),
            ("reviewing", "리뷰 관련 작업 유형", "백로그 내 일반 작업으로 생성될 수 있음"),
            ("testing", "테스트 작업 유형", "일반 작업으로 처리"),
            ("deploying", "리뷰 통과 후 배포", "Deployment Frequency, Lead Time 완료"),
            ("incident", "결함/장애 대응", "CFR, Recovery Time, interruption, energy 비용"),
        ],
        [1600, 3500, 4260],
    )
    add_table(
        doc,
        ["난이도", "요구 skill", "base steps", "defect probability", "요구 domain knowledge", "생성 비중"],
        [
            ("C1", "0.5", "1", "0.05", "0.20", "20%"),
            ("C2", "1.0", "2", "0.10", "0.35", "35%"),
            ("C3", "1.5", "3", "0.18", "0.50", "25%"),
            ("C4", "2.0", "5", "0.28", "0.65", "15%"),
            ("C5", "2.5", "8", "0.40", "0.80", "5%"),
        ],
        [1200, 1500, 1400, 1900, 2200, 1160],
        font_size=8.4,
    )


def section_flow(doc):
    doc.add_heading("4. 실행 흐름과 관계식", level=1)
    doc.add_heading("4.1 매 step 실행 순서", level=2)
    add_table(
        doc,
        ["순서", "처리 내용"],
        [
            ("1", "Step 카운터 증가 및 PM 개입 capacity 초기화"),
            ("2", "스프린트 시작 조건이면 PL sprint planning 실행"),
            ("3", "PM 요구사항 조율 이벤트 시도"),
            ("4", "환경 이벤트로 incident 발생 여부 평가"),
            ("5", "team_tech_debt 소폭 증가 및 backlog pressure 적용"),
            ("6", "meeting/interruption pressure 적용"),
            ("7", "DeveloperAgent와 PLAgent를 임의 순서로 step 실행"),
            ("8", "PM 병목 감지 및 완화 개입 시도"),
            ("9", "deploying task 처리 및 지표 기록"),
            ("10", "스프린트 종료 조건이면 sprint retro 실행"),
        ],
        [1000, 8360],
        font_size=8.6,
    )

    doc.add_heading("4.2 핵심 관계식", level=2)
    add_table(
        doc,
        ["관계", "계산/규칙", "해석"],
        [
            ("Backlog pressure", "raw = max(0, (backlog/active_devs - 5.0) / 5.0)", "개발자당 작업량이 목표 WIP 5개를 넘으면 energy/motivation 하락"),
            ("PL awareness 완화", "effective = raw * (1 - avg_team_awareness * 0.4)", "PL이 팀 상태를 잘 파악할수록 백로그 압박이 일부 완화"),
            ("Meeting pressure", "max(0, (meeting_load - 60) / 60)", "일일 회의가 60분을 넘는 부분만 추가 압박으로 작동"),
            ("Review cost", "1 + 0.8 * review_strictness", "리뷰 엄격도가 높을수록 energy 소모와 리뷰 소요가 증가"),
            ("Review defect reduction", "defect_prob * (1 - 0.6 * review_strictness)", "엄격한 리뷰는 결함 유출 가능성을 낮춤"),
            ("Effective clarity", "requirement_clarity + 0.35 * requirement_coordination", "PM 요구사항 조율이 켜진 경우 요구사항 명확도 상승"),
            ("Incident probability", "(1 - stability) * 0.15 * review/clarity/PM multipliers", "코드 안정성, 리뷰, 요구사항, PM 조율이 결함 발생에 함께 작용"),
            ("Help request", "knowledge_gap * help_seeking_tendency", "업무 요구 도메인 지식보다 낮으면 도움 요청 가능성 증가"),
            ("Help success", "task progress +0.05, requester domain knowledge +0.01, helper energy -1.5", "멘토링은 진행률과 학습을 높이지만 helper 부담을 만든다"),
        ],
        [1900, 3650, 3810],
        font_size=8.0,
    )

    add_callout(
        doc,
        "관계 해석 원칙",
        "이 모델은 특정 파라미터가 하나의 지표만 바꾸는 구조가 아닙니다. 예를 들어 review_strictness는 CFR을 낮추는 동시에 Lead Time과 energy 비용을 높일 수 있습니다. "
        "따라서 실험 결과는 단일 지표 최적화가 아니라 품질, 속도, 에너지, 백로그의 trade-off로 읽어야 합니다.",
    )


def section_initial_settings(doc):
    doc.add_heading("5. 초기 세팅과 기본 분포", level=1)
    doc.add_heading("5.1 모델 기본값", level=2)
    add_table(
        doc,
        ["파라미터", "기본값", "의미", "주요 사용 위치"],
        [
            ("num_developers", "9", "DeveloperAgent 수", "model/simulate/dashboard"),
            ("num_pl", "1", "PLAgent 수", "model"),
            ("num_sprints", "6", "실행 스프린트 수", "model/simulate/dashboard"),
            ("meeting_load", "60.0", "일일 회의 부하 분", "meeting pressure"),
            ("review_strictness", "0.7", "리뷰 엄격도", "review cost/defect"),
            ("codebase_stability", "0.8", "코드베이스 안정성", "incident probability"),
            ("tech_debt_ratio", "0.1", "초기 팀 기술부채", "team_tech_debt"),
            ("pipeline_efficiency", "0.7", "파이프라인 효율 변수", "현재 보관 중심"),
            ("requirement_clarity", "0.6", "요구사항 명확도", "defect/incident probability"),
            ("communication_overhead", "0.3", "커뮤니케이션 오버헤드", "현재 보관 중심"),
            ("knowledge_decay_rate", "0.02", "지식 감쇠율", "Developer knowledge decay"),
            ("collaboration_tendency", "0.6", "지식 전파 강도", "transfer_knowledge"),
            ("sprint_backlog_size", "30", "스프린트당 작업 후보 수", "backlog generation/planning"),
            ("seed", "42", "난수 시드", "재현성"),
        ],
        [1750, 1150, 3850, 2610],
        font_size=8.0,
    )

    doc.add_heading("5.2 역할과 팀 구성", level=2)
    add_table(
        doc,
        ["team_composition", "junior", "middle", "senior", "설명"],
        [
            ("junior_heavy", "5", "3", "1", "초급자 비중이 높은 팀"),
            ("balanced", "3", "4", "2", "균형형 팀"),
            ("senior_heavy", "1", "4", "4", "시니어 비중이 높은 팀"),
            ("None", "skill distribution", "skill distribution", "skill distribution", "명시적 role pool 없이 기본 skill 분포 사용"),
        ],
        [2100, 1100, 1100, 1100, 3960],
        font_size=8.4,
    )

    doc.add_heading("5.3 초기 분포", level=2)
    add_table(
        doc,
        ["변수", "분포", "범위/값", "설명"],
        [
            ("developer.skill_level", "lognormal", "0.5~3.0, median=base skill", "역할/skill pool 주변에서 개인 역량 변동"),
            ("developer.learning_rate", "beta(5,5)", "0.1~0.9", "중간값 중심 학습률"),
            ("developer.burnout_threshold", "truncated normal", "mean=20, sd=4, 10~30", "번아웃 임계값"),
            ("developer.quality_tendency", "lognormal", "median=0.6, 0.3~0.95", "품질 성향"),
            ("developer.interrupt_sensitivity", "beta(3,5)", "0.1~0.8", "회의/중단 민감도"),
            ("developer.energy", "truncated normal", "mean=85, sd=8, 0~100", "초기 에너지"),
            ("developer.motivation", "truncated normal", "mean=70, sd=12, 0~100", "초기 동기"),
            ("developer.knowledge", "truncated normal", "mean=skill*20, sd=12, 0~100", "초기 코드/업무 지식"),
            ("pl.team_awareness", "uniform", "0.6~0.9", "PL의 팀 상태 감지 역량"),
        ],
        [2300, 2050, 2450, 2560],
        font_size=8.0,
    )

    add_callout(
        doc,
        "초기값 보정 방법",
        "실측 데이터가 확보되면 sampling.py의 DEFAULT_DISTRIBUTIONS를 직접 수정하거나, 실험 조건에서 distribution_overrides를 넘겨 특정 분포를 상수 또는 다른 분포로 교체할 수 있습니다.",
        fill="F4F6F9",
    )


def section_usage(doc):
    doc.add_heading("6. 사용 방법", level=1)
    doc.add_heading("6.1 환경 준비", level=2)
    add_para(doc, "프로젝트 폴더 기준 위치:")
    add_code(doc, "cd LG_CNS_VNB_ABM/abm_ver2")
    add_para(doc, "필수 Python 패키지는 requirements.txt에 정의되어 있습니다.")
    add_code(doc, "pip install -r requirements.txt")

    doc.add_heading("6.2 웹 UI 실행", level=2)
    add_code(doc, "python server.py")
    add_para(doc, "브라우저에서 다음 주소를 엽니다.")
    add_code(doc, "http://localhost:8765")
    add_table(
        doc,
        ["UI 슬라이더", "범위", "서버 전달 파라미터", "의미"],
        [
            ("개발자 수", "3~20", "num_developers", "팀 크기"),
            ("스프린트 수", "2~12", "num_sprints", "시뮬레이션 기간"),
            ("회의 부하", "0~240, 15분 단위", "meeting_load", "일일 회의 시간"),
            ("리뷰 엄격도", "0.1~1.0", "review_strictness", "품질/속도 trade-off"),
            ("코드베이스 안정성", "0.1~1.0", "codebase_stability", "incident 발생 기반"),
            ("협업 성향", "0.1~1.0", "collaboration_tendency", "지식 전파 강도"),
            ("요구사항 명확도", "0.1~1.0", "requirement_clarity", "결함/재작업 위험"),
            ("랜덤 시드", "1~999", "seed", "난수 재현성"),
        ],
        [1750, 1700, 2300, 3610],
        font_size=8.3,
    )

    doc.add_heading("6.3 단일 시뮬레이션 실행", level=2)
    add_para(doc, "웹 UI 없이 JSON 파라미터를 넘겨 단일 실행 결과를 얻을 수 있습니다.")
    add_code(doc, "python simulate.py '{\"num_sprints\": 6, \"meeting_load\": 120, \"review_strictness\": 0.7}'")
    add_para(doc, "반환값은 params, snapshots, prism, internal_metrics, total_steps, num_sprints로 구성됩니다.")

    doc.add_heading("6.4 반복 실험 실행", level=2)
    add_code(doc, "python experiment_runner.py --scenario A --runs 30 --sprints 6 --seed-start 1000 --output-dir outputs/scenario_A")
    add_para(doc, "A~E 중 하나의 시나리오를 선택할 수 있으며, 결과는 run-level CSV와 summary CSV로 저장됩니다.")
    add_code(doc, "python run_final_experiments.py")
    add_para(doc, "최종 실험 스크립트는 A~E 시나리오를 각각 100회 반복하고 outputs/final에 결과를 보존합니다.")


def section_params(doc):
    doc.add_heading("7. 실험 파라미터 카탈로그", level=1)
    doc.add_heading("7.1 사용자가 직접 조작하기 좋은 파라미터", level=2)
    add_table(
        doc,
        ["파라미터", "권장 범위/선택지", "주요 효과", "대표 시나리오"],
        [
            ("meeting_load", "30, 60, 120, 180 또는 0~240", "회의가 baseline 60분을 넘으면 energy/motivation/flow에 압박", "A"),
            ("requirement_clarity", "0.4, 0.8 또는 0.1~1.0", "높을수록 결함/incident 위험 감소", "A/E"),
            ("review_strictness", "0.3, 0.5, 0.7, 0.9", "높을수록 CFR 감소, 리뷰 비용/Lead Time 증가 가능", "B"),
            ("codebase_stability", "0.4, 0.8 또는 0.1~1.0", "높을수록 incident 기본 발생률 감소", "B"),
            ("sprint_backlog_size", "10, 20, 30, 40, 50", "큰 백로그는 backlog pressure와 잔여 작업 증가", "C"),
            ("pl.team_awareness", "0.4, 0.8 등 override", "높을수록 번아웃 감지와 압박 완화에 유리", "C"),
            ("team_composition", "junior_heavy, balanced, senior_heavy", "역할 구성에 따른 도움 요청/멘토링/생산성 변화", "D"),
            ("mentoring_intensity", "0.3, 0.8", "높을수록 도움 해결 가능성 증가, helper 부담도 증가", "D"),
            ("pm_profile", "weak/allocation/bottleneck/requirement/strong", "PM 개입 능력 조합", "E"),
            ("allocation_skill", "0.3 또는 0.8", "domain-fit 업무 배정 가능성 증가", "E"),
            ("bottleneck_detection", "0.3 또는 0.8", "병목 감지/재배정/멘토 과부하 완화 가능성 증가", "E"),
            ("requirement_coordination", "0.3 또는 0.8", "effective requirement clarity 상승, CFR 감소", "E"),
        ],
        [1900, 2100, 3800, 1560],
        font_size=7.8,
    )

    doc.add_heading("7.2 시나리오 구성", level=2)
    add_table(
        doc,
        ["시나리오", "실험 질문", "조건"],
        [
            ("A", "회의 부하와 요구사항 명확도는 생산성과 품질에 어떤 영향을 주는가?", "requirement_clarity 0.4/0.8 x meeting_load 30/60/120/180"),
            ("B", "리뷰 엄격도는 품질과 속도 사이에서 어떤 trade-off를 만드는가?", "codebase_stability 0.4/0.8 x review_strictness 0.3/0.5/0.7/0.9"),
            ("C", "백로그 크기와 PL team awareness는 에너지와 처리량을 어떻게 바꾸는가?", "team_awareness 0.4/0.8 x sprint_backlog_size 10~50"),
            ("D", "팀 구성과 멘토링 강도는 지식 전파와 시니어 부담에 어떤 영향을 주는가?", "team_composition 3종 x mentoring_intensity 0.3/0.8"),
            ("E", "PM의 배정/병목/요구사항 조율 역량은 생산성, 품질, 에너지에 어떤 효과를 내는가?", "weak, allocation-focused, bottleneck-focused, requirement-focused, strong PM"),
        ],
        [950, 4750, 3660],
        font_size=8.0,
    )

    doc.add_heading("7.3 PM 프로파일", level=2)
    add_table(
        doc,
        ["pm_profile", "allocation", "bottleneck", "req_coord", "해석"],
        [
            ("weak_pm", "0.3", "0.3", "0.3", "전반적 PM 개입 역량 낮음"),
            ("allocation_focused_pm", "0.8", "0.3", "0.3", "업무 배정 최적화 중심"),
            ("bottleneck_focused_pm", "0.3", "0.8", "0.3", "병목 감지/완화 중심"),
            ("requirement_focused_pm", "0.3", "0.3", "0.8", "요구사항 조율 중심"),
            ("strong_pm", "0.8", "0.8", "0.8", "모든 PM 개입 역량 높음"),
        ],
        [2300, 1600, 1900, 2100, 1460],
        font_size=8.0,
    )


def section_modify(doc):
    doc.add_heading("8. 수정 방법과 확장 체크리스트", level=1)
    add_callout(
        doc,
        "수정 원칙",
        "파라미터 하나를 추가할 때는 model.py에서 저장하고, simulate.py에서 입력받고, experiment_runner.py의 PARAM_COLUMNS와 결과 CSV에 반영하며, "
        "웹에서 조작해야 한다면 dashboard.html의 슬라이더와 runSimulation() 파라미터에도 추가해야 합니다.",
        fill="F4F6F9",
    )
    add_table(
        doc,
        ["수정 목적", "수정 파일", "확인할 항목"],
        [
            ("새 정책 변수 추가", "model.py, simulate.py, experiment_runner.py, dashboard.html", "기본값, 범위, CSV 컬럼, UI 전달값"),
            ("새 시나리오 추가", "experiment_runner.py, run_final_experiments.py, make_report_assets.py", "condition 함수, choices, 최종 루프, 보고용 표/그림"),
            ("업무 난이도/분포 변경", "tasks.py", "COMPLEXITY_CONFIG, COMPLEXITY_DIST, domain requirement"),
            ("도메인 추가", "tasks.py, agents.py, simulate.py", "TASK_DOMAINS, domain_knowledge 초기화, coverage 계산"),
            ("개발자 행동 변경", "agents.py", "state 전이, productivity, energy/motivation/knowledge 변화"),
            ("PL 행동 변경", "agents.py", "standup, coaching, assignment, incident handling, retro"),
            ("전역 압박/관계식 변경", "model.py", "backlog pressure, meeting pressure, incident probability, PM 개입"),
            ("새 출력 지표 추가", "simulate.py, experiment_runner.py, dashboard.html", "internal_metrics, CSV columns, 화면 표시"),
            ("초기 분포 보정", "sampling.py 또는 distribution_overrides", "분포 타입, low/high, seed 재현성"),
        ],
        [1900, 3100, 4360],
        font_size=8.0,
    )

    doc.add_heading("8.1 새 파라미터 추가 절차", level=2)
    for item in [
        "model.py의 LGCNSDevModel.__init__ 인자에 파라미터와 기본값을 추가합니다.",
        "self.<parameter> 형태로 모델 상태에 저장하고, 필요한 관계식에서 사용합니다.",
        "simulate.py의 run_simulation(params)에서 params.get('<parameter>', default)를 넘깁니다.",
        "반복 실험 대상이면 experiment_runner.py의 PARAM_COLUMNS에 추가하고 scenario 조건에 값을 넣습니다.",
        "웹 UI에서 조작할 값이면 dashboard.html에 입력 컨트롤을 추가하고 runSimulation()의 params 객체에 넣습니다.",
        "결과 해석이 필요하면 internal_metrics 또는 PRISM 출력에 반영하고, 요약 CSV 컬럼을 확인합니다.",
        "동일 seed에서 기존 시나리오가 깨지지 않는지 smoke run을 수행합니다.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("8.2 새 시나리오 추가 절차", level=2)
    add_code(doc, "def scenario_f_conditions(): ...")
    add_para(doc, "새 조건 함수를 만든 뒤 get_conditions(), argparse choices, run_final_experiments.py의 SCENARIOS에 추가합니다.")
    add_para(doc, "보고서용 표와 그림까지 필요하면 make_report_assets.py에 scenario_F_summary.csv를 읽는 로직과 표/그림 생성 함수를 추가합니다.")


def section_outputs(doc):
    doc.add_heading("9. 산출물 해석과 공유 시 유의사항", level=1)
    add_table(
        doc,
        ["산출물", "위치", "내용"],
        [
            ("experiment_results.csv", "outputs 또는 지정 output-dir", "각 run별 파라미터, PRISM 지표, 내부 지표"),
            ("experiment_summary.csv", "outputs 또는 지정 output-dir", "condition별 mean/std 요약"),
            ("scenario_*_results.csv", "outputs/final", "최종 A~E run-level 결과 보존본"),
            ("scenario_*_summary.csv", "outputs/final", "최종 A~E condition 요약 보존본"),
            ("report_tables/*.csv", "outputs/final/report_tables", "보고서 삽입용 정리 표"),
            ("figures/*.png", "outputs/final/figures", "보고서 삽입용 그래프"),
        ],
        [2300, 2600, 4460],
        font_size=8.3,
    )
    doc.add_heading("공유 시 해석 가이드", level=2)
    add_bullet(doc, "단일 run보다 condition별 반복 평균과 표준편차를 우선 해석합니다.")
    add_bullet(doc, "PRISM 6개 지표는 대외 공유에 적합하고, internal_metrics는 원인 진단과 모델 디버깅에 적합합니다.")
    add_bullet(doc, "시나리오 비교는 baseline 대비 증감률과 함께 해석해야 합니다.")
    add_bullet(doc, "품질 개선, 속도 개선, 에너지 유지, 잔여 백로그 감소는 동시에 최적화되지 않을 수 있습니다.")
    add_bullet(doc, "실측 캘리브레이션 전에는 절대값 예측보다 방향성, 민감도, trade-off 비교로 사용합니다.")

    doc.add_heading("현재 알려진 주의사항", level=2)
    add_table(
        doc,
        ["항목", "주의 내용", "권장 조치"],
        [
            ("요구사항 명확도 기본값", "코드 기본값은 0.6이나 웹 UI 초기 표시값은 0.70입니다.", "실험 보고 시 실행 경로와 입력값을 명시"),
            ("UI 노출 범위", "웹 UI는 D/E 시나리오 파라미터를 직접 노출하지 않습니다.", "고급 실험은 experiment_runner.py 또는 API 사용"),
            ("실측 데이터", "분포와 계수 상당 부분은 추정값입니다.", "Git/Jira/HR/DXI 데이터로 보정"),
            ("난수성", "seed와 반복 횟수에 따라 결과가 흔들릴 수 있습니다.", "충분한 runs와 seed_start 고정"),
            ("배포 설정", "Docker는 requirements.txt를 설치하고 server.py를 실행합니다.", "배포 환경 PORT 설정 확인"),
        ],
        [1800, 4200, 3360],
        font_size=8.0,
    )


def section_appendix(doc):
    doc.add_heading("부록. 빠른 참조 표", level=1)
    doc.add_heading("A. 조작 경로별 파라미터", level=2)
    add_table(
        doc,
        ["조작 경로", "사용 가능 항목"],
        [
            ("웹 UI", "num_developers, num_sprints, meeting_load, review_strictness, codebase_stability, collaboration_tendency, requirement_clarity, seed"),
            ("simulate.py/API", "웹 UI 항목 + sprint_backlog_size, team_composition, mentoring_intensity, pm_profile, allocation_skill, bottleneck_detection, requirement_coordination, pm_intervention_capacity, distribution_overrides"),
            ("experiment_runner.py", "A~E 사전 정의 조건. PARAM_COLUMNS에 포함된 실험 변수 중심"),
            ("sampling.py", "developer/pl 초기 분포 자체를 보정"),
            ("model.py/agents.py/tasks.py", "관계식, 상태 전이, 업무 분포, 에너지/지식/품질 로직 자체를 수정"),
        ],
        [2100, 7260],
        font_size=8.0,
    )
    doc.add_heading("B. 내부 지표 묶음", level=2)
    add_table(
        doc,
        ["묶음", "지표"],
        [
            ("상태", "avg_energy, avg_motivation, min_energy, avg_knowledge, low_energy_count"),
            ("도메인", "avg_team_domain_knowledge, min_team_domain_knowledge, domain_coverage, domain_mismatch_count"),
            ("멘토링", "help_requests_total, help_requests_resolved, help_request_resolution_rate, mentoring_load_total, helper_interruptions"),
            ("역할", "junior_count, middle_count, senior_count, junior_avg_knowledge, senior_mentoring_load, junior_help_requests"),
            ("PM", "allocation_match_score, bottlenecks_detected, bottleneck_interventions, reassignments, clarification_events, effective_requirement_clarity"),
            ("운영", "attrition_count, coaching_count, remaining_backlog, completed_tasks, active_developers"),
        ],
        [1800, 7560],
        font_size=8.0,
    )


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_manual_toc(doc)
    section_overview(doc)
    section_structure(doc)
    section_design(doc)
    section_flow(doc)
    section_initial_settings(doc)
    section_usage(doc)
    section_params(doc)
    section_modify(doc)
    section_outputs(doc)
    section_appendix(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
