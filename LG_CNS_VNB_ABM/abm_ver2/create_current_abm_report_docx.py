#!/usr/bin/env python3
from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
FINAL_DIR = BASE_DIR / "outputs" / "final"
TABLE_DIR = FINAL_DIR / "report_tables"
FIGURE_DIR = FINAL_DIR / "figures"
FGHI_RUN_ROOT = BASE_DIR / "outputs" / "validation_runs" / "20260626_165229_calibrated"
FGHI_REPORT_DIR = BASE_DIR / "outputs" / "reports" / "fghi_calibrated_report_20260629_211737"
FGHI_FIGURE_DIR = FGHI_REPORT_DIR / "figures"
OUT_PATH = BASE_DIR / "PRISM_ABM_현행구조_및_A-I_시나리오_실험결과_보고서.docx"

FONT = "맑은 고딕"
TEXT = RGBColor(32, 36, 44)
MUTED = RGBColor(91, 101, 116)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
TABLE_BORDER = "CBD3DF"
WHITE = RGBColor(255, 255, 255)

CONTENT_WIDTH_DXA = 9720


def read_rows(name: str) -> list[dict]:
    path = TABLE_DIR / name
    with path.open("r", newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


ROWS = {
    "A": read_rows("scenario_A_report_table.csv"),
    "B": read_rows("scenario_B_report_table.csv"),
    "C": read_rows("scenario_C_report_table.csv"),
    "D": read_rows("scenario_D_report_table.csv"),
    "E": read_rows("scenario_E_report_table.csv"),
}


def read_fghi_summary(scenario: str) -> list[dict]:
    path = FGHI_RUN_ROOT / f"scenario_{scenario}" / "experiment_summary.csv"
    with path.open("r", newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


FGHI_ROWS = {
    "F": read_fghi_summary("F"),
    "G": read_fghi_summary("G"),
    "H": read_fghi_summary("H"),
    "I": read_fghi_summary("I"),
}


def by_condition(scenario: str) -> dict[str, dict]:
    return {row["condition_id"]: row for row in ROWS[scenario]}


def fval(row: dict, key: str) -> float:
    try:
        return float(row.get(key, "") or 0.0)
    except (TypeError, ValueError):
        return 0.0


def mean_group_rows(scenario: str, group_key: str, order: list[str], metrics: list[tuple[str, str, int]]):
    grouped = {}
    for row in FGHI_ROWS[scenario]:
        grouped.setdefault(row.get(group_key, ""), []).append(row)

    output = []
    keys = order or sorted(grouped)
    for key in keys:
        rows = grouped.get(str(key), [])
        if not rows:
            continue
        record = [str(key).replace("_", " ")]
        for _, metric_key, digits in metrics:
            value = sum(fval(row, metric_key) for row in rows) / len(rows)
            record.append(f"{value:.{digits}f}")
        output.append(record)
    return output


def set_run_font(run, size=None, color=None, bold=None, italic=None, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size=None, color=None, bold=None):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=TABLE_BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_header_repeat(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.875)
    section.bottom_margin = Inches(0.875)
    section.left_margin = Inches(0.875)
    section.right_margin = Inches(0.875)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=10.5, color=TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for level, size, color, before, after in [
        (1, 16, BLUE, 15, 7),
        (2, 13, BLUE, 10, 5),
        (3, 11.5, DARK_BLUE, 7, 3),
    ]:
        style = styles[f"Heading {level}"]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_style = styles["List Bullet"]
    set_style_font(list_style, size=10.2, color=TEXT)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.12

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("PRISM ABM 구조 및 실험 결과 보고서")
    set_run_font(r, size=8.5, color=MUTED)


def add_para(doc, text="", style=None, size=10.5, color=TEXT, bold=False, italic=False,
             align=None, before=0, after=5, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3.5)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=10.2, color=TEXT)
    return p


def add_caption(doc, text):
    return add_para(doc, text, size=8.8, color=MUTED, italic=True, after=7)


def add_callout(doc, title, bullets, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA], indent_dxa=0)
    set_table_borders(table, color="D5DCE6", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    for text in bullets:
        bp = cell.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        br = bp.add_run(text)
        set_run_font(br, size=9.8, color=TEXT)
    add_para(doc, "", after=5)


def format_cell_text(cell, size=8.8, bold=False, color=TEXT, align=None):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        if align is not None:
            p.alignment = align
        for run in p.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_header_repeat(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        format_cell_text(cell, size=8.8, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(str(value)) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
            format_cell_text(cells[i], size=font_size, color=TEXT, align=align)
    add_para(doc, "", after=5)
    return table


def v(row, key, digits=1):
    raw = row.get(key, "")
    try:
        number = float(raw)
    except (TypeError, ValueError):
        return raw
    if digits == 0:
        return f"{number:.0f}"
    return f"{number:.{digits}f}"


def pct(value):
    try:
        return f"{float(value):+.1f}%"
    except (TypeError, ValueError):
        return str(value)


def add_cover(doc: Document):
    add_para(doc, "PRISM ABM", size=24, color=DARK_BLUE, bold=True, after=1)
    add_para(doc, "Agent-Based Modeling 구조 및 시나리오 실험 결과 보고", size=15, color=TEXT, bold=True, after=13)
    add_para(doc, "현행 구현 기준 요약본", size=11, color=MUTED, after=22)

    rows = [
        ("대상 코드", "LG_CNS_VNB_ABM/abm_ver2"),
        ("근거 산출물", "outputs/final, calibrated validation run, F~I report figures"),
        ("실험 범위", "Scenario A~I, 조건별 100회 반복"),
        ("시뮬레이션 기준", "9 Developer Agent, 1 PL Agent, 6 Sprint"),
        ("작성 기준일", "2026.06.30"),
    ]
    add_table(doc, ["구분", "내용"], rows, [2100, 7620], font_size=9.3)
    add_callout(doc, "핵심 요약", [
        "모델은 개발자 상태, 업무 흐름, PL/PM 개입, 환경 압박의 상호작용을 시뮬레이션한다.",
        "회의 부하와 백로그 압박은 평균 에너지와 개인 처리량을 크게 낮춘다.",
        "요구사항 명확도, 리뷰 엄격도, PM 요구사항 조율은 실패율을 낮추는 품질 레버로 작동한다.",
        "멘토링은 도움 해결률을 높이지만 helper interruption과 mentoring load를 함께 증가시킨다.",
        "프로젝트 유형, 팀 구성, PM profile, 요구사항 변동성은 F~I 시나리오에서 추가 검증한다.",
    ])
    doc.add_page_break()


def section_overview(doc: Document):
    doc.add_heading("1. 보고서 개요", level=1)
    add_bullet(doc, "본 보고서는 현재 디렉토리에 구현된 PRISM ABM의 구조와 A~I 시나리오 실험 결과를 요약한다.")
    add_bullet(doc, "기존 보고서 양식을 참고하여 개조식 문장과 표 중심으로 작성한다.")
    add_bullet(doc, "모델 설명은 현재 코드의 DeveloperAgent, PLAgent, Task, Environment 정의를 기준으로 한다.")
    add_bullet(doc, "A~E 결과는 outputs/final의 summary CSV, report table, figure 파일을 기준으로 한다.")
    add_bullet(doc, "F~I 결과는 calibrated validation run의 summary CSV와 F~I 보고용 figure 파일을 기준으로 한다.")
    add_bullet(doc, "해석은 단일 실행값보다 조건별 100회 반복 평균과 조건 간 변화율을 우선한다.")

    doc.add_heading("1.1 문서 구성", level=2)
    add_table(doc, ["장", "내용"], [
        ("2장", "모델링 시스템의 구성 요소와 관계를 설명한다."),
        ("3장", "에이전트, 태스크, 환경 변수, 산출 지표를 정리한다."),
        ("4장", "A~I 시나리오 실험 목적과 공통 세팅을 설명한다."),
        ("5장", "A~E 시나리오 핵심 결과와 시각화 자료를 제시한다."),
        ("6장", "추가 F~I 시나리오 결과를 제시한다."),
        ("7장", "종합 시사점과 해석 유의사항을 정리한다."),
    ], [1100, 8620], font_size=9.2)
    doc.add_page_break()


def section_model_system(doc: Document):
    doc.add_heading("2. 모델링 시스템 상세 설명", level=1)
    add_bullet(doc, "본 모델은 소프트웨어 개발팀을 작은 가상 조직으로 표현하는 Mesa 기반 ABM이다.")
    add_bullet(doc, "모델의 목적은 정책 변화가 생산성, 품질, 피로도, 협업 부하에 주는 방향성을 비교하는 것이다.")
    add_bullet(doc, "모델은 Developer Agent, PL Agent, Task, Environment, PM Intervention으로 구성된다.")
    add_bullet(doc, "1 step은 1일이며 1 sprint는 10 step이다.")
    add_bullet(doc, "기본 실험은 9명의 Developer Agent, 1명의 PL Agent, 6 sprint, 총 60 step으로 실행한다.")
    add_bullet(doc, "각 조건은 난수 seed를 바꾸어 100회 반복하고 평균값으로 해석한다.")

    add_callout(doc, "모델을 읽는 기본 관점", [
        "Developer Agent는 에너지, 동기, 지식, 역량이 높을수록 일을 빠르게 처리한다.",
        "PL Agent는 업무 배정, 리뷰 배정, 코칭, 인시던트 대응을 통해 팀 운영을 담당한다.",
        "Task는 난이도, 도메인, 결함 가능성, 진행률을 가진 업무 단위이다.",
        "Environment는 회의 부하, 요구사항 명확도, 리뷰 정책, 백로그 압박을 만든다.",
        "PM Intervention은 배정 최적화, 병목 완화, 요구사항 조율, 범위 통제를 수행한다.",
    ])

    doc.add_heading("2.1 시스템 구성 요소", level=2)
    add_table(doc, ["구성 요소", "모델 내 의미", "주요 속성", "대표 출력"], [
        ("Developer Agent", "개발자 개인이다.", "skill, role, energy, motivation, knowledge, domain_knowledge이다.", "PR, 리뷰, 도움 요청, 멘토링, 이탈이다."),
        ("PL Agent", "팀 리더이다.", "team_awareness, leadership_style, energy, coaching_count이다.", "업무 배정, 리뷰 배정, 코칭, 회고 조정이다."),
        ("Task", "처리해야 할 업무이다.", "task_type, complexity, domain, progress, defect_prob이다.", "완료, 배포, 재작업, 인시던트이다."),
        ("Environment", "프로젝트 조건이다.", "meeting_load, backlog, review_strictness, requirement_clarity이다.", "에너지 소모, 실패율, 리드타임, 잔여 백로그이다."),
        ("PM Intervention", "PM 역량 개입이다.", "allocation, bottleneck, requirement, scope control이다.", "배정 적합도, 명확도, 범위 변경, 병목 완화이다."),
    ], [1500, 1950, 3500, 2770], font_size=8.0)

    doc.add_heading("2.2 실행 단위와 시간 구조", level=2)
    add_table(doc, ["단위", "정의", "모델 내 처리"], [
        ("Step", "1일이다.", "환경 이벤트, PM 개입, Agent 행동, 배포, 지표 기록이 순서대로 실행된다."),
        ("Sprint", "10 step이다.", "시작 시 planning을 수행하고 종료 시 velocity와 retro를 수행한다."),
        ("Run", "하나의 seed로 수행한 60 step 실행이다.", "한 조건의 단일 관측값이다."),
        ("Condition", "시나리오의 한 실험 조합이다.", "예: meeting_load 180, requirement_clarity 0.8이다."),
        ("Scenario", "조건 집합이다.", "A~I가 각각 다른 정책 질문을 검증한다."),
    ], [1350, 2600, 5770], font_size=8.5)

    doc.add_heading("2.3 초기화 흐름", level=2)
    add_bullet(doc, "모델은 입력 파라미터와 seed를 받아 LGCNSDevModel 객체를 생성한다.")
    add_bullet(doc, "project_type이 있으면 요구사항 명확도, 코드 안정성, 백로그 크기, 업무 난이도 분포를 기본값으로 설정한다.")
    add_bullet(doc, "team_composition이 있으면 junior, middle, senior 인원 구성을 먼저 만든다.")
    add_bullet(doc, "team_composition이 없으면 기본 skill 분포로 9명의 개발자를 만든다.")
    add_bullet(doc, "각 개발자는 역량, 에너지, 동기, 지식, 도메인 지식, 협업 성향을 가진다.")
    add_bullet(doc, "PL Agent는 전체 Developer Agent를 team_members로 가진다.")
    add_bullet(doc, "초기 backlog는 sprint_backlog_size × num_sprints 개수로 생성한다.")
    add_bullet(doc, "Task는 생성 시 task_type, complexity, domain, defect_prob, required_skill을 가진다.")

    doc.add_page_break()

    doc.add_heading("2.4 1일 Step 처리 순서", level=2)
    add_table(doc, ["순서", "처리 내용", "효과"], [
        ("1", "current_step을 1 증가시킨다.", "시뮬레이션 날짜가 이동한다."),
        ("2", "스프린트 첫날이면 planning을 수행한다.", "PL과 개발자 에너지가 일부 소모되고 리뷰 대기가 배정된다."),
        ("3", "PM 요구사항 조율을 시도한다.", "effective_requirement_clarity가 품질 위험을 낮춘다."),
        ("4", "요구사항 변동성과 scope change를 평가한다.", "진행률 rollback, 복잡도 상승, 재작업, 신규 backlog가 발생할 수 있다."),
        ("5", "인시던트 발생 여부를 평가한다.", "코드 안정성, 리뷰 엄격도, 명확도, PM 조율이 확률에 반영된다."),
        ("6", "백로그 압박과 회의 압박을 적용한다.", "개발자 에너지, 동기, flow_streak가 감소할 수 있다."),
        ("7", "PL과 Developer Agent가 각자 행동한다.", "업무 배정, 코딩, 리뷰, 도움 요청, 코칭이 수행된다."),
        ("8", "PM 병목 완화 개입을 시도한다.", "저성과 업무 재배정 또는 helper 회복이 발생할 수 있다."),
        ("9", "배포 가능한 Task를 배포 처리한다.", "deployments, lead_time, new capability 지표가 갱신된다."),
        ("10", "지표를 기록하고 스프린트 종료 여부를 확인한다.", "시계열과 sprint velocity가 저장된다."),
    ], [700, 3550, 5470], font_size=8.0)

    doc.add_heading("2.5 업무 생애주기", level=2)
    add_table(doc, ["단계", "상태", "설명"], [
        ("생성", "backlog", "업무가 난이도, 도메인, 유형을 가진 backlog 항목으로 생성된다."),
        ("배정", "in_progress", "PL이 유휴 개발자에게 업무를 배정한다."),
        ("개발", "Coding 또는 FlowCoding", "개발자의 생산성만큼 progress가 증가한다."),
        ("PR", "review_pending", "progress가 1.0 이상이면 PR이 생성되고 리뷰 대기열로 이동한다."),
        ("리뷰", "Reviewing", "리뷰어가 review_strictness와 review_capacity에 따라 검토한다."),
        ("결과", "deploying 또는 backlog", "정상 업무는 배포 단계로 가고, 결함/요구사항 문제는 재작업으로 돌아간다."),
        ("배포", "done", "배포가 성공하면 lead_time과 deployment가 기록된다."),
        ("장애", "incident", "실패 또는 환경 이벤트로 인시던트가 발생하며 복구 시간이 기록된다."),
    ], [850, 1850, 7020], font_size=8.4)

    doc.add_heading("2.6 Developer Agent 생산성 구조", level=2)
    add_bullet(doc, "개발자 생산성은 skill_coeff × energy/100 × motivation/100 × knowledge/100으로 계산한다.")
    add_bullet(doc, "skill_coeff는 skill_level이 높을수록 커지는 역량 계수이다.")
    add_bullet(doc, "energy는 회의, 리뷰, 코딩, 인시던트, 멘토링, 백로그 압박으로 감소한다.")
    add_bullet(doc, "motivation은 시간이 지나며 조금씩 감소하고 코칭, 완료 경험, 회고 보너스로 증가한다.")
    add_bullet(doc, "knowledge는 시간이 지나며 감소하고 학습, 리뷰, 멘토링 성공으로 증가한다.")
    add_bullet(doc, "세 요소 중 하나가 낮아지면 전체 생산성이 곱셈 구조로 함께 낮아진다.")
    add_bullet(doc, "이 구조는 역량이 높아도 에너지 또는 지식이 낮으면 처리량이 낮아지는 상황을 표현한다.")

    doc.add_heading("2.7 에이전트 간 관계", level=2)
    add_table(doc, ["관계", "발생 조건", "모델 내 영향"], [
        ("PL -> Developer", "유휴 개발자와 backlog가 존재한다.", "task_assignment가 발생하고 도메인 매칭 점수가 기록된다."),
        ("Developer -> Task", "개발자가 업무를 보유한다.", "progress가 증가하고 PR 또는 재작업으로 전이된다."),
        ("Developer -> Developer", "도메인 지식 gap이 존재한다.", "help_request, help_success, mentoring event가 기록된다."),
        ("Reviewer -> Task", "review_pending 업무가 존재한다.", "결함, 재작업, 배포 전이가 결정된다."),
        ("PM -> Team", "PM profile 또는 PM 역량 변수가 설정된다.", "allocation, bottleneck, requirement, scope intervention이 발생한다."),
        ("Environment -> Team", "회의, 백로그, 변동성, 인시던트 조건이 존재한다.", "에너지, 동기, flow, 품질 위험이 변한다."),
    ], [1600, 3000, 5120], font_size=8.1)

    doc.add_page_break()

    doc.add_heading("2.8 PM/PL 개입 구조", level=2)
    add_table(doc, ["개입", "수행 주체", "작동 방식", "확인 지표"], [
        ("Task allocation", "PL 또는 PM", "도메인 지식, 역량, 에너지, 부하를 보고 업무를 배정한다.", "allocation_match_score, domain_mismatch_count이다."),
        ("Burnout coaching", "PL", "에너지가 낮은 개발자를 team_awareness 확률로 감지해 회복시킨다.", "coaching_count, avg_energy이다."),
        ("Review assignment", "PL", "review_capacity, 도메인 지식, skill, energy가 높은 리뷰어를 선택한다.", "review_assignment_count, CFR이다."),
        ("Requirement coordination", "PM", "요구사항 명확도를 보정하고 결함 확률을 낮춘다.", "effective_requirement_clarity, CFR이다."),
        ("Bottleneck intervention", "PM", "저진행 업무를 재배정하거나 과부하 helper를 회복시킨다.", "bottlenecks_detected, reassignments이다."),
        ("Scope control", "PM", "요구사항 변동에 따른 범위 변경을 예방한다.", "scope_changes_prevented, rework_rate이다."),
    ], [1450, 1050, 4700, 2520], font_size=7.8)

    doc.add_heading("2.9 주요 입력값 사용법", level=2)
    add_table(doc, ["입력값", "권장 해석", "높아질 때의 일반 효과"], [
        ("meeting_load", "팀의 회의·동기화 부담이다.", "에너지와 몰입이 낮아지고 PR/Engineer가 감소한다."),
        ("requirement_clarity", "요구사항의 명확성이다.", "결함, 재작업, 실패율이 감소한다."),
        ("review_strictness", "리뷰 기준의 엄격도이다.", "결함 유출은 줄지만 리뷰 비용이 증가한다."),
        ("codebase_stability", "코드베이스 안정성이다.", "인시던트 발생 위험과 CFR이 낮아진다."),
        ("sprint_backlog_size", "스프린트 업무량이다.", "백로그 압박과 잔여 업무가 증가한다."),
        ("team_awareness", "PL의 팀 상태 인지 수준이다.", "코칭과 압박 완충 효과가 커진다."),
        ("mentoring_intensity", "멘토링 활성 수준이다.", "도움 해결률과 senior/helper 부하가 함께 증가한다."),
        ("requirement_volatility", "요구사항 변경 빈도이다.", "scope change, rollback, rework가 증가한다."),
        ("pm_profile", "PM 역량 조합이다.", "배정, 병목, 요구사항, 범위 통제 효과가 달라진다."),
    ], [2150, 3000, 4570], font_size=8.0)

    doc.add_heading("2.10 모델 사용 흐름", level=2)
    add_bullet(doc, "먼저 비교하려는 정책 질문을 정한다.")
    add_bullet(doc, "다음으로 시나리오 A~I 중 질문에 가장 가까운 실험 구조를 선택한다.")
    add_bullet(doc, "조건별 조작 변수와 반복 횟수, seed 범위를 정한다.")
    add_bullet(doc, "experiment_runner는 run-level 결과와 condition-level summary를 생성한다.")
    add_bullet(doc, "PRISM 지표는 최종 성과를 비교할 때 사용한다.")
    add_bullet(doc, "내부 지표는 왜 그런 결과가 나왔는지 원인을 해석할 때 사용한다.")
    add_bullet(doc, "개별 숫자보다 조건 간 차이와 방향성을 우선 해석한다.")
    add_bullet(doc, "실측 데이터가 연결되면 초기 분포와 계수를 보정해 같은 구조로 다시 실행한다.")


def section_agents_tasks(doc: Document):
    doc.add_heading("3. 에이전트, 업무, 지표 정의", level=1)
    doc.add_heading("3.1 Developer Agent 속성", level=2)
    add_table(doc, ["속성군", "주요 변수", "초기화 기준", "행동 영향"], [
        ("역량/역할", "skill_level, role", "skill 또는 team_composition으로 설정한다.", "생산성, 리뷰 능력, 멘토링 성향을 결정한다."),
        ("상태", "energy, motivation, knowledge", "기본값은 에너지 100, 동기 70, 지식은 skill 기반이다.", "생산성 공식에 직접 곱해진다."),
        ("몰입", "flow_streak, state", "초기 상태는 Idle이다.", "연속 코딩 시 FlowCoding으로 전환된다."),
        ("품질", "quality_tendency", "skill과 경험 기반으로 설정한다.", "기술부채와 결함 가능성에 영향을 준다."),
        ("중단 민감도", "interrupt_sensitivity", "개발자별로 다르게 부여한다.", "회의 압박에 의한 flow disruption을 키운다."),
        ("도메인 지식", "domain_knowledge", "6개 도메인별로 생성한다.", "도움 요청, 배정 적합도, 재작업 위험에 영향을 준다."),
        ("협업", "prior_collaboration_score", "개발자별 협업 이력으로 설정한다.", "도움 성공률과 커뮤니케이션 비용에 영향을 준다."),
        ("멘토링", "help_seeking, mentoring_capacity", "역량과 도메인 지식으로 설정한다.", "도움 요청 빈도와 helper 부하를 만든다."),
    ], [1350, 2300, 2950, 3120], font_size=7.8)

    doc.add_heading("3.2 Developer Agent 상태 전이", level=2)
    add_table(doc, ["상태", "의미", "전이 조건"], [
        ("Idle", "업무가 없는 대기 상태이다.", "업무를 받으면 Coding으로 이동하고, 확률적으로 Learning으로 이동한다."),
        ("Coding", "업무 progress를 올리는 상태이다.", "progress 완료 시 review_pending으로 이동한다."),
        ("FlowCoding", "연속 코딩으로 몰입한 상태이다.", "flow_streak가 3 이상이고 energy가 60 초과일 때 진입한다."),
        ("Reviewing", "다른 개발자의 PR을 검토하는 상태이다.", "review_progress가 1.0 이상이면 결과를 결정한다."),
        ("Learning", "대기 중 학습 상태이다.", "knowledge가 증가하고 다시 Coding 또는 Idle로 돌아간다."),
        ("Interrupted", "인시던트 또는 외부 중단 상태이다.", "energy가 감소하고 flow_streak가 초기화된다."),
        ("Burnout", "에너지가 임계값 이하인 상태이다.", "5 step 지속 시 attrition으로 처리된다."),
    ], [1450, 3550, 4720], font_size=8.1)

    doc.add_heading("3.3 역할과 팀 구성", level=2)
    add_table(doc, ["팀 구성", "Junior", "Middle", "Senior", "해석"], [
        ("junior_heavy", "5명", "3명", "1명", "학습과 멘토링 수요가 큰 팀이다."),
        ("balanced", "3명", "4명", "2명", "일반적인 혼합 팀이다."),
        ("senior_heavy", "1명", "4명", "4명", "처리 안정성과 도메인 대응력이 높은 팀이다."),
    ], [1900, 1050, 1050, 1050, 4670], font_size=8.4)
    add_bullet(doc, "role이 지정되지 않으면 skill_level로 junior, middle, senior를 추정한다.")
    add_bullet(doc, "junior는 학습률과 도움 요청 성향이 높고 멘토링 제공 성향은 낮다.")
    add_bullet(doc, "senior는 도메인 지식과 멘토링 제공 성향이 높고 도움 요청 성향은 낮다.")
    add_bullet(doc, "역할은 staffing_cost, senior mentoring load, junior help requests 해석에 직접 사용된다.")

    doc.add_heading("3.4 PL Agent 행동", level=2)
    add_table(doc, ["행동", "실행 시점", "세부 설명"], [
        ("Daily Standup", "매 step", "PL과 개발자 에너지를 소모하고 개발자의 flow_streak를 초기화한다."),
        ("Burnout Coaching", "매 step", "에너지가 낮은 개발자를 team_awareness 확률로 감지해 에너지와 동기를 회복시킨다."),
        ("Incident Assignment", "인시던트 발생 후", "Critical은 전체 팀을 중단시키고 High는 담당자를 중단시킨다."),
        ("Task Assignment", "유휴 개발자 존재 시", "skill과 도메인 지식을 고려해 backlog task를 배정한다."),
        ("Review Assignment", "스프린트 planning", "review_capacity와 도메인 지식이 높은 개발자를 리뷰어로 선택한다."),
        ("Sprint Retro", "10 step마다", "평균 에너지와 동기를 보고 다음 sprint backlog와 review strictness를 조정한다."),
    ], [1800, 1850, 6070], font_size=8.1)

    doc.add_page_break()

    doc.add_heading("3.5 Task 정의", level=2)
    add_table(doc, ["항목", "값", "설명"], [
        ("task_type", "coding, reviewing, testing, deploying, incident", "업무 유형이며 흐름과 지표 집계 경로를 결정한다."),
        ("complexity", "C1, C2, C3, C4, C5", "난이도이며 required_skill, base_steps, defect_prob를 결정한다."),
        ("domain", "frontend, backend, data, infra, legacy, testing", "업무 도메인이며 도메인 지식과 배정 적합도에 연결된다."),
        ("progress", "0.0~1.0", "개발 진행률이며 1.0 이상이면 PR이 생성된다."),
        ("status", "backlog, in_progress, review_pending, done", "업무 생애주기 상태이다."),
        ("rework_count", "0 이상", "요구사항, 리뷰, 테스트, 도메인 mismatch 등으로 되돌아간 횟수이다."),
    ], [1550, 3100, 5070], font_size=8.0)

    doc.add_heading("3.6 Complexity 정의", level=2)
    add_table(doc, ["난이도", "required_skill", "base_steps", "defect_prob", "domain knowledge 요구"], [
        ("C1", "0.5", "1", "0.05", "0.20"),
        ("C2", "1.0", "2", "0.10", "0.35"),
        ("C3", "1.5", "3", "0.18", "0.50"),
        ("C4", "2.0", "5", "0.28", "0.65"),
        ("C5", "2.5", "8", "0.40", "0.80"),
    ], [1200, 1850, 1550, 1550, 3570], font_size=8.4)
    add_bullet(doc, "skill_level이 required_skill보다 낮으면 effective_steps가 증가한다.")
    add_bullet(doc, "complexity가 높을수록 소요 step과 defect_prob가 증가한다.")
    add_bullet(doc, "도메인 지식이 요구 수준보다 낮으면 help_request와 domain_mismatch 위험이 커진다.")

    doc.add_heading("3.7 Project Type 기본값", level=2)
    add_table(doc, ["Project Type", "특성", "모델 기본값 해석"], [
        ("new_build", "신규 구축이다.", "요구사항 명확도와 코드 안정성이 중간 수준이며 신규 기능 비중이 높다."),
        ("maintenance_enhancement", "운영 개선이다.", "요구사항과 코드 안정성이 비교적 높고 legacy 업무가 포함된다."),
        ("legacy_migration", "레거시 전환이다.", "코드 안정성이 낮고 고난이도와 legacy 도메인 비중이 높다."),
        ("deadline_driven", "납기 압박 프로젝트이다.", "백로그 크기와 요구사항 변동성이 높다."),
        ("quality_critical", "품질 중요 프로젝트이다.", "리뷰·테스트 비중과 품질 gate 비용이 높고 결함 유출이 낮다."),
    ], [2300, 2300, 5120], font_size=8.1)

    doc.add_heading("3.8 PM Profile 정의", level=2)
    add_table(doc, ["PM Profile", "allocation", "bottleneck", "requirement", "scope", "해석"], [
        ("weak_pm", "0.3", "0.3", "0.3", "0.3", "모든 개입 역량이 낮다."),
        ("allocation_focused_pm", "0.8", "0.3", "0.3", "0.4", "업무 배정 최적화에 강하다."),
        ("bottleneck_focused_pm", "0.3", "0.8", "0.3", "0.4", "병목 감지와 완화에 강하다."),
        ("requirement_focused_pm", "0.3", "0.3", "0.8", "0.6", "요구사항 조율에 강하다."),
        ("strong_pm", "0.8", "0.8", "0.8", "0.8", "전 영역 개입 역량이 높다."),
    ], [2300, 1150, 1250, 1350, 1000, 2670], font_size=7.8)

    doc.add_heading("3.9 산출 지표", level=2)
    add_table(doc, ["지표군", "지표", "계산/해석"], [
        ("PRISM", "PRs per Engineer", "활성 개발자당 PR 생성량이다."),
        ("PRISM", "Lead Time (steps)", "업무 생성부터 배포까지 평균 소요 step이다."),
        ("PRISM", "Deployment Frequency", "10 step 기준 스프린트당 배포 빈도이다."),
        ("PRISM", "Change Failure Rate (%)", "failed_deployments / deployments × 100이다."),
        ("PRISM", "Recovery Time (steps)", "인시던트 생성부터 복구 완료까지 평균 step이다."),
        ("PRISM", "% Time on New Capabilities", "신규 기능 관련 투입 비중이다."),
        ("Internal", "avg_energy, avg_motivation, avg_knowledge", "팀 상태를 해석하는 내부 진단 지표이다."),
        ("Internal", "remaining_backlog, completed_tasks", "처리량과 미처리 업무 압박을 해석한다."),
        ("Internal", "help_resolution_rate, mentoring_load", "멘토링 효과와 helper 부하를 해석한다."),
        ("Internal", "allocation_match_score, domain_mismatch_count", "PM/PL 배정 품질을 해석한다."),
        ("Internal", "rework_rate, scope_changes", "요구사항·범위 변경의 품질 비용을 해석한다."),
    ], [1350, 3050, 5320], font_size=7.9)


def section_experiment_design(doc: Document):
    doc.add_heading("4. 시나리오 실험 설계", level=1)
    add_bullet(doc, "본 보고서는 A~E 최종 실험과 F~I calibrated validation 실험을 함께 기준으로 한다.")
    add_bullet(doc, "모든 시나리오는 9명의 Developer Agent와 1명의 PL Agent로 실행한다.")
    add_bullet(doc, "모든 시나리오는 6 sprint, 60 step 기준으로 실행한다.")
    add_bullet(doc, "각 condition은 seed_start 1000부터 100회 반복 실행한다.")
    add_bullet(doc, "결과는 run-level CSV와 condition-level summary CSV로 저장한다.")

    add_table(doc, ["시나리오", "실험 목적", "조건 수", "주요 조작 변수"], [
        ("A", "회의 부하와 요구사항 명확도의 영향을 확인한다.", "8", "meeting_load, requirement_clarity이다."),
        ("B", "리뷰 엄격도와 코드베이스 안정성의 품질 효과를 확인한다.", "8", "review_strictness, codebase_stability이다."),
        ("C", "백로그 압박과 팀 인지 수준의 효과를 확인한다.", "10", "sprint_backlog_size, team_awareness이다."),
        ("D", "팀 구성과 멘토링 강도의 효과를 확인한다.", "6", "team_composition, mentoring_intensity이다."),
        ("E", "PM 역량 프로파일별 개입 효과를 확인한다.", "5", "pm_profile이다."),
        ("F", "프로젝트 유형별 적합한 팀 구성을 확인한다.", "15", "project_type, team_composition이다."),
        ("G", "프로젝트 유형별 PM profile 효과를 확인한다.", "25", "project_type, pm_profile이다."),
        ("H", "팀 구성, 멘토링 강도, 프로젝트 유형의 상호작용을 확인한다.", "27", "project_type, team_composition, mentoring_intensity이다."),
        ("I", "요구사항 변동성과 PM 요구사항 조율 효과를 확인한다.", "27", "project_type, requirement_volatility, pm_profile이다."),
    ], [900, 3920, 800, 4100], font_size=7.8)

    doc.add_heading("4.1 전체 민감도 요약", level=2)
    add_bullet(doc, "아래 그림은 A~C 주요 조건 변화가 핵심 지표에 미치는 방향을 한 화면에 요약한다.")
    add_picture(doc, FIGURE_DIR / "scenario_impact_heatmap.png", width=Inches(5.7))
    add_caption(doc, "그림 1. Scenario A~C 주요 변화율 Heatmap이다.")


def add_picture(doc, path: Path, width=Inches(5.9)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=width)
    return p


def section_scenario_a(doc: Document):
    doc.add_heading("5.1 Scenario A: 회의 부하와 요구사항 명확도", level=2)
    add_bullet(doc, "실험 목적은 meeting_load 증가가 생산성과 에너지에 미치는 영향을 확인하는 것이다.")
    add_bullet(doc, "추가 목적은 requirement_clarity 개선이 실패율을 낮추는지 확인하는 것이다.")
    add_bullet(doc, "실험 조건은 meeting_load 30, 60, 120, 180과 requirement_clarity 0.4, 0.8의 조합이다.")

    a = by_condition("A")
    rows = []
    for cid in ["A2", "A4", "A6", "A8"]:
        r = a[cid]
        rows.append([
            cid,
            v(r, "meeting_load", 0),
            v(r, "requirement_clarity", 1),
            v(r, "PRs per Engineer mean", 2),
            v(r, "Change Failure Rate (%) mean", 2),
            v(r, "avg_energy mean", 1),
            v(r, "coaching_count mean", 1),
        ])
    add_table(doc, ["조건", "회의", "명확도", "PR/Eng", "CFR", "Energy", "Coaching"], rows,
              [850, 850, 900, 1100, 1100, 1150, 1500], font_size=8.4)

    add_bullet(doc, "회의 부하가 60에서 180으로 증가하면 PR/Engineer는 2.21에서 1.61로 27.3% 감소한다.")
    add_bullet(doc, "같은 조건에서 평균 에너지는 62.3에서 44.2로 29.0% 감소한다.")
    add_bullet(doc, "같은 조건에서 코칭 횟수는 5.95에서 49.28로 크게 증가한다.")
    add_bullet(doc, "요구사항 명확도가 0.4에서 0.8로 상승하면 CFR은 4.50에서 3.75로 16.8% 감소한다.")
    add_bullet(doc, "회의 부하는 생산성 저하 레버이며 요구사항 명확도는 품질 개선 레버로 해석된다.")
    add_picture(doc, FIGURE_DIR / "scenario_A_meeting_vs_prs.png", width=Inches(5.45))
    add_caption(doc, "그림 2. Scenario A meeting_load별 PR/Engineer 변화이다.")


def section_scenario_b(doc: Document):
    doc.add_heading("5.2 Scenario B: 리뷰 엄격도와 코드베이스 안정성", level=2)
    add_bullet(doc, "실험 목적은 review_strictness가 Change Failure Rate를 낮추는지 확인하는 것이다.")
    add_bullet(doc, "추가 목적은 codebase_stability가 품질 기준선을 어떻게 바꾸는지 확인하는 것이다.")
    add_bullet(doc, "실험 조건은 review_strictness 0.3, 0.5, 0.7, 0.9와 codebase_stability 0.4, 0.8의 조합이다.")

    b = by_condition("B")
    rows = []
    for cid in ["B1", "B4", "B5", "B8"]:
        r = b[cid]
        rows.append([
            cid,
            v(r, "review_strictness", 1),
            v(r, "codebase_stability", 1),
            v(r, "PRs per Engineer mean", 2),
            v(r, "Lead Time (steps) mean", 2),
            v(r, "Change Failure Rate (%) mean", 2),
            v(r, "avg_energy mean", 1),
        ])
    add_table(doc, ["조건", "리뷰", "안정성", "PR/Eng", "Lead", "CFR", "Energy"], rows,
              [850, 850, 950, 1100, 1050, 1150, 1250], font_size=8.4)

    add_bullet(doc, "불안정 코드베이스에서 리뷰 엄격도 0.3 대비 0.9는 CFR을 16.62에서 8.88로 46.6% 낮춘다.")
    add_bullet(doc, "안정적 코드베이스에서도 같은 리뷰 강화는 CFR을 5.97에서 3.92로 34.4% 낮춘다.")
    add_bullet(doc, "코드베이스 안정성 0.8은 안정성 0.4보다 전반적인 실패율 기준선이 낮다.")
    add_bullet(doc, "현재 결과에서는 리뷰 강화가 리드타임을 크게 악화시키지 않는다.")
    add_bullet(doc, "리뷰 엄격도는 품질 안정화에 직접적인 효과를 갖는 정책 변수로 해석된다.")
    add_picture(doc, FIGURE_DIR / "scenario_B_quality_speed_tradeoff.png", width=Inches(5.45))
    add_caption(doc, "그림 3. Scenario B 품질-속도 Trade-off 요약이다.")


def section_scenario_c(doc: Document):
    doc.add_heading("5.3 Scenario C: 백로그 압박과 팀 인지 수준", level=2)
    add_bullet(doc, "실험 목적은 sprint_backlog_size 증가가 처리량, 에너지, 잔여 백로그에 미치는 영향을 확인하는 것이다.")
    add_bullet(doc, "추가 목적은 team_awareness가 백로그 압박을 완화하는지 확인하는 것이다.")
    add_bullet(doc, "실험 조건은 backlog size 10, 20, 30, 40, 50과 team_awareness 0.4, 0.8의 조합이다.")

    c = by_condition("C")
    rows = []
    for cid in ["C1", "C5", "C6", "C10"]:
        r = c[cid]
        rows.append([
            cid,
            v(r, "sprint_backlog_size", 0),
            v(r, "team_awareness", 1),
            v(r, "PRs per Engineer mean", 2),
            v(r, "avg_energy mean", 1),
            v(r, "remaining_backlog mean", 1),
            v(r, "coaching_count mean", 1),
        ])
    add_table(doc, ["조건", "Backlog", "Awareness", "PR/Eng", "Energy", "잔여", "Coaching"], rows,
              [850, 1050, 1200, 1100, 1150, 1300, 1300], font_size=8.4)

    add_bullet(doc, "team_awareness 0.4에서 backlog 10 대비 50은 PR/Engineer를 2.72에서 2.08로 23.6% 낮춘다.")
    add_bullet(doc, "같은 조건에서 평균 에너지는 84.4에서 53.7로 36.4% 감소한다.")
    add_bullet(doc, "같은 조건에서 잔여 백로그는 22.0에서 242.1로 약 10배 증가한다.")
    add_bullet(doc, "team_awareness 0.8 조건에서도 backlog 50의 에너지 감소와 잔여 백로그 증가는 크게 남는다.")
    add_bullet(doc, "팀 인지와 코칭은 완충 장치이지만 과도한 업무량 자체를 대체하지 못한다.")
    add_picture(doc, FIGURE_DIR / "scenario_C_backlog_vs_remaining_backlog.png", width=Inches(5.45))
    add_caption(doc, "그림 4. Scenario C backlog size별 잔여 백로그 변화이다.")


def section_scenario_d(doc: Document):
    doc.add_heading("5.4 Scenario D: 팀 구성과 멘토링 강도", level=2)
    add_bullet(doc, "실험 목적은 junior/middle/senior 구성과 mentoring_intensity가 지원 네트워크와 생산성에 미치는 영향을 확인하는 것이다.")
    add_bullet(doc, "실험 조건은 junior_heavy, balanced, senior_heavy 팀 구성과 mentoring_intensity 0.3, 0.8의 조합이다.")

    rows = []
    for r in ROWS["D"]:
        rows.append([
            r["condition_id"],
            r["team_composition"],
            v(r, "mentoring_intensity", 1),
            v(r, "PRs per Engineer mean", 2),
            v(r, "help_request_resolution_rate mean", 2),
            v(r, "mentoring_load_total mean", 1),
            v(r, "helper_interruptions mean", 1),
        ])
    add_table(doc, ["조건", "팀 구성", "멘토링", "PR/Eng", "해결률", "Load", "Interrupt"], rows,
              [750, 1850, 900, 1050, 1050, 1050, 1200], font_size=8.1)

    add_bullet(doc, "junior_heavy에서 mentoring_intensity 0.3 대비 0.8은 도움 해결률을 0.29에서 0.80으로 높인다.")
    add_bullet(doc, "balanced에서도 도움 해결률은 0.30에서 0.75로 높아진다.")
    add_bullet(doc, "senior_heavy는 high mentoring 조건에서 PR/Engineer 3.36으로 가장 높은 처리량을 보인다.")
    add_bullet(doc, "junior_heavy high mentoring은 mentoring load를 9.0에서 21.5로 크게 높인다.")
    add_bullet(doc, "멘토링은 지식 전달 효과와 helper interruption 비용을 함께 만든다.")
    add_picture(doc, FIGURE_DIR / "scenario_D_mentoring_load_tradeoff.png", width=Inches(6.0))
    add_caption(doc, "그림 5. Scenario D 멘토링 강도와 부하 Trade-off이다.")


def section_scenario_e(doc: Document):
    doc.add_heading("5.5 Scenario E: PM 프로파일별 개입 효과", level=2)
    add_bullet(doc, "실험 목적은 PM 역량 프로파일별로 업무 배분, 병목 완화, 요구사항 조율 효과를 확인하는 것이다.")
    add_bullet(doc, "실험 조건은 weak, allocation-focused, bottleneck-focused, requirement-focused, strong PM의 5개 프로파일이다.")

    rows = []
    for r in ROWS["E"]:
        rows.append([
            r["condition_id"],
            r["pm_profile"].replace("_", " "),
            v(r, "PRs per Engineer mean", 2),
            v(r, "avg_energy mean", 1),
            v(r, "Change Failure Rate (%) mean", 2),
            v(r, "allocation_match_score mean", 2),
            v(r, "domain_mismatch_count mean", 1),
            v(r, "bottleneck_interventions mean", 1),
        ])
    add_table(doc, ["조건", "PM", "PR/Eng", "Energy", "CFR", "Match", "Mismatch", "Interv."], rows,
              [700, 2250, 950, 1000, 950, 900, 1000, 1050], font_size=7.9)

    add_bullet(doc, "allocation-focused PM은 allocation match를 0.71에서 0.77로 높이고 domain mismatch를 60.3% 낮춘다.")
    add_bullet(doc, "bottleneck-focused PM은 평균 에너지를 65.2에서 69.1로 6.0% 높인다.")
    add_bullet(doc, "bottleneck-focused PM은 병목 개입과 재배정을 크게 늘려 처리량을 5.1% 높인다.")
    add_bullet(doc, "requirement-focused PM은 CFR을 3.54에서 2.16으로 38.9% 낮춘다.")
    add_bullet(doc, "strong PM은 에너지, 품질, 도메인 매칭을 균형 있게 개선한다.")
    add_bullet(doc, "단일 역량 강화는 목표 지표에는 효과적이나 모든 지표를 동시에 개선하지는 않는다.")
    add_picture(doc, FIGURE_DIR / "scenario_E_pm_impact_summary.png", width=Inches(6.0))
    add_caption(doc, "그림 6. Scenario E PM 프로파일별 핵심 지표 변화율이다.")


def section_results(doc: Document):
    doc.add_heading("5. 시나리오별 실험 결과", level=1)
    add_bullet(doc, "각 시나리오는 목적, 세팅, 대표 결과, 해석 순서로 정리한다.")
    add_bullet(doc, "표의 수치는 조건별 100회 반복 평균이다.")
    add_bullet(doc, "CFR은 Change Failure Rate를 의미한다.")
    section_scenario_a(doc)
    section_scenario_b(doc)
    section_scenario_c(doc)
    section_scenario_d(doc)
    section_scenario_e(doc)


def section_scenario_f(doc: Document):
    doc.add_heading("6.1 Scenario F: 프로젝트 유형과 팀 구성", level=2)
    add_bullet(doc, "실험 목적은 프로젝트 유형별로 적합한 팀 구성이 달라지는지 확인하는 것이다.")
    add_bullet(doc, "실험 조건은 5개 project_type과 3개 team_composition의 조합이다.")
    add_bullet(doc, "프로젝트 유형은 new_build, maintenance_enhancement, legacy_migration, deadline_driven, quality_critical이다.")

    rows = mean_group_rows("F", "team_composition", ["junior_heavy", "balanced", "senior_heavy"], [
        ("PR/Eng", "PRs per Engineer mean", 2),
        ("PR/Cost", "PRs_per_cost mean", 2),
        ("완료/Cost", "completed_tasks_per_cost mean", 2),
        ("Rework", "rework_rate mean", 2),
        ("잔여", "remaining_backlog mean", 1),
        ("Senior Load", "senior_mentoring_load_per_senior mean", 1),
    ])
    add_table(doc, ["팀 구성", "PR/Eng", "PR/Cost", "완료/Cost", "Rework", "잔여", "Senior Load"], rows,
              [1700, 950, 1000, 1150, 950, 1150, 1450], font_size=8.1)

    add_bullet(doc, "시니어 중심 팀은 평균 PR/Engineer 2.95로 가장 높다.")
    add_bullet(doc, "시니어 중심 팀은 잔여 백로그 137.9로 가장 낮다.")
    add_bullet(doc, "주니어 중심 팀은 완료 업무/Cost 1.61로 비용 대비 완료량이 가장 높다.")
    add_bullet(doc, "주니어 중심 팀은 Senior Load 31.2로 시니어 1인당 멘토링 부담이 가장 높다.")
    add_bullet(doc, "팀 구성 평가는 처리량, 비용 효율, 시니어 부하를 함께 보아야 한다.")
    add_picture(doc, FGHI_FIGURE_DIR / "scenario_F_team_composition_dashboard.png", width=Inches(6.05))
    add_caption(doc, "그림 7. Scenario F 프로젝트 유형과 팀 구성별 성과 대시보드이다.")


def section_scenario_g(doc: Document):
    doc.add_page_break()
    doc.add_heading("6.2 Scenario G: 프로젝트 유형과 PM Profile", level=2)
    add_bullet(doc, "실험 목적은 프로젝트 유형별로 PM profile이 품질, 범위관리, 에너지에 미치는 효과를 확인하는 것이다.")
    add_bullet(doc, "실험 조건은 5개 project_type과 5개 pm_profile의 조합이다.")
    add_bullet(doc, "PM profile은 weak, allocation-focused, bottleneck-focused, requirement-focused, strong PM이다.")

    rows = mean_group_rows("G", "pm_profile", [
        "weak_pm",
        "allocation_focused_pm",
        "bottleneck_focused_pm",
        "requirement_focused_pm",
        "strong_pm",
    ], [
        ("CFR", "Change Failure Rate (%) mean", 2),
        ("Rework", "rework_rate mean", 2),
        ("Scope", "scope_changes mean", 2),
        ("Prevent", "scope_changes_prevented mean", 2),
        ("Energy", "avg_energy mean", 1),
        ("PM Cap", "pm_capacity_used mean", 1),
    ])
    add_table(doc, ["PM", "CFR", "Rework", "Scope", "Prevent", "Energy", "PM Cap"], rows,
              [2250, 900, 950, 950, 1050, 1000, 1050], font_size=7.7)

    add_bullet(doc, "strong PM은 scope changes prevented 3.10으로 가장 높다.")
    add_bullet(doc, "strong PM은 rework rate 0.61로 가장 낮다.")
    add_bullet(doc, "requirement-focused PM은 CFR 4.20%로 가장 낮다.")
    add_bullet(doc, "strong PM은 PM capacity used 39.7로 가장 높아 개입 비용도 함께 증가한다.")
    add_bullet(doc, "PM profile은 품질, 범위관리, 운영 에너지에 서로 다른 방식으로 작동한다.")
    add_picture(doc, FGHI_FIGURE_DIR / "scenario_G_pm_profile_dashboard.png", width=Inches(6.05))
    add_caption(doc, "그림 8. Scenario G 프로젝트 유형과 PM profile별 성과 대시보드이다.")


def section_scenario_h(doc: Document):
    doc.add_page_break()
    doc.add_heading("6.3 Scenario H: 팀 구성, 멘토링 강도, 프로젝트 유형", level=2)
    add_bullet(doc, "실험 목적은 멘토링 강도가 도움 해결률, 시니어 부하, 병목, 재작업에 미치는 영향을 확인하는 것이다.")
    add_bullet(doc, "실험 조건은 3개 project_type, 3개 team_composition, 3개 mentoring_intensity의 조합이다.")
    add_bullet(doc, "프로젝트 유형은 new_build, legacy_migration, deadline_driven이다.")

    rows = mean_group_rows("H", "mentoring_intensity", ["0.2", "0.5", "0.8"], [
        ("Help", "help_resolution_rate mean", 2),
        ("Senior Load", "senior_mentoring_load_per_senior mean", 1),
        ("Bottleneck", "senior_bottleneck_index mean", 2),
        ("Rework", "rework_rate mean", 2),
        ("Energy", "avg_energy mean", 1),
    ])
    add_table(doc, ["멘토링", "Help", "Senior Load", "Bottleneck", "Rework", "Energy"], rows,
              [1150, 950, 1450, 1350, 1050, 1100], font_size=8.1)

    add_bullet(doc, "멘토링 강도 0.2에서 0.8로 높이면 help resolution rate는 0.15에서 0.59로 상승한다.")
    add_bullet(doc, "같은 구간에서 senior mentoring load per senior는 4.23에서 13.31로 상승한다.")
    add_bullet(doc, "주니어 중심 팀은 senior mentoring load per senior 19.84로 가장 높다.")
    add_bullet(doc, "멘토링 강화는 도움 해결률을 높이지만 시니어 부하를 함께 키운다.")
    add_bullet(doc, "재작업률은 멘토링 강도만으로 뚜렷하게 낮아지지 않는다.")
    add_picture(doc, FGHI_FIGURE_DIR / "scenario_H_mentoring_help_load_dashboard.png", width=Inches(6.05))
    add_caption(doc, "그림 9. Scenario H 멘토링 강도별 도움 해결률과 시니어 부하이다.")


def section_scenario_i(doc: Document):
    doc.add_page_break()
    doc.add_heading("6.4 Scenario I: 요구사항 변동성과 PM Requirement Coordination", level=2)
    add_bullet(doc, "실험 목적은 requirement_volatility가 rework, scope change, lead time, CFR에 미치는 영향을 확인하는 것이다.")
    add_bullet(doc, "실험 조건은 3개 project_type, 3개 requirement_volatility, 3개 pm_profile의 조합이다.")
    add_bullet(doc, "PM profile은 weak, requirement-focused, strong PM이다.")

    rows = mean_group_rows("I", "requirement_volatility", ["0.2", "0.5", "0.8"], [
        ("Rework", "rework_rate mean", 2),
        ("Scope", "scope_changes mean", 2),
        ("Prevent", "scope_changes_prevented mean", 2),
        ("Lead", "Lead Time (steps) mean", 2),
        ("CFR", "Change Failure Rate (%) mean", 2),
        ("Energy", "avg_energy mean", 1),
    ])
    add_table(doc, ["변동성", "Rework", "Scope", "Prevent", "Lead", "CFR", "Energy"], rows,
              [1050, 950, 950, 1050, 950, 950, 1050], font_size=8.1)

    rows_pm = mean_group_rows("I", "pm_profile", ["weak_pm", "requirement_focused_pm", "strong_pm"], [
        ("Rework", "rework_rate mean", 2),
        ("Scope", "scope_changes mean", 2),
        ("Prevent", "scope_changes_prevented mean", 2),
        ("CFR", "Change Failure Rate (%) mean", 2),
        ("Energy", "avg_energy mean", 1),
    ])
    add_table(doc, ["PM", "Rework", "Scope", "Prevent", "CFR", "Energy"], rows_pm,
              [2300, 1000, 1000, 1100, 1000, 1100], font_size=8.0)

    add_bullet(doc, "요구사항 변동성 0.2에서 0.8로 높이면 rework rate는 0.59에서 3.28로 증가한다.")
    add_bullet(doc, "같은 구간에서 scope changes는 1.76에서 3.40으로 증가한다.")
    add_bullet(doc, "strong PM은 scope changes prevented 4.40으로 가장 높고 rework rate 1.45로 가장 낮다.")
    add_bullet(doc, "requirement-focused PM은 CFR 2.85%로 가장 낮다.")
    add_bullet(doc, "요구사항 변동성은 재작업과 범위 변경을 키우며, strong PM은 이를 가장 강하게 완충한다.")
    add_picture(doc, FGHI_FIGURE_DIR / "scenario_I_volatility_rework_scope_dashboard.png", width=Inches(6.05))
    add_caption(doc, "그림 10. Scenario I 요구사항 변동성별 재작업과 범위 변경이다.")


def section_additional_results(doc: Document):
    doc.add_page_break()
    doc.add_heading("6. 추가 시나리오 F~I 실험 결과", level=1)
    add_bullet(doc, "F~I 시나리오는 calibrated validation run의 summary CSV를 기준으로 작성한다.")
    add_bullet(doc, "모든 수치는 조건별 100회 반복 평균이다.")
    add_bullet(doc, "F~I는 프로젝트 유형, 팀 구성, PM profile, 요구사항 변동성을 더 세분화해 확인한다.")
    section_scenario_f(doc)
    section_scenario_g(doc)
    section_scenario_h(doc)
    section_scenario_i(doc)


def section_implications(doc: Document):
    doc.add_heading("7. 종합 시사점", level=1)
    add_bullet(doc, "회의 부하와 백로그 압박은 개발자 에너지를 낮추고 처리량을 줄이는 핵심 압박 요인이다.")
    add_bullet(doc, "요구사항 명확도와 PM 요구사항 조율은 실패율을 낮추는 품질 개선 요인이다.")
    add_bullet(doc, "리뷰 엄격도는 코드베이스가 불안정할수록 더 큰 품질 개선 효과를 보인다.")
    add_bullet(doc, "팀 구성은 생산성 수준을 크게 바꾸는 구조적 변수이다.")
    add_bullet(doc, "멘토링은 junior-heavy 팀에서 도움 해결률을 크게 높이지만 senior/helper 부하를 함께 키운다.")
    add_bullet(doc, "PM 개입은 목표 역량에 따라 효과가 다르게 나타난다.")
    add_bullet(doc, "시니어 중심 팀은 처리량과 잔여 백로그 관점에서 안정적이나 비용 효율은 별도 검토가 필요하다.")
    add_bullet(doc, "요구사항 변동성 증가는 rework와 scope change를 크게 키우며 strong PM이 가장 강하게 완충한다.")
    add_bullet(doc, "실험 결과는 절대 예측값보다 정책 변화의 방향성과 trade-off 탐색에 적합하다.")

    doc.add_heading("7.1 보고 활용 시 유의사항", level=2)
    add_bullet(doc, "현재 결과는 실측 캘리브레이션 이전의 구조 검증 및 민감도 분석 결과이다.")
    add_bullet(doc, "Git, Jira, CI/CD, Incident, HR 데이터가 연결되면 초기 분포와 계수를 보정해야 한다.")
    add_bullet(doc, "BA와 Bridge Manager는 별도 Agent가 아니라 requirement_clarity와 communication_overhead로 흡수되어 있다.")
    add_bullet(doc, "단일 시나리오의 최적 조건보다 품질, 속도, 에너지, 잔여 백로그 간 균형을 함께 봐야 한다.")

    add_callout(doc, "운영 관점 결론", [
        "단기 생산성 개선은 회의 부하와 백로그 압박 관리에서 시작한다.",
        "품질 안정화는 요구사항 명확도, 리뷰 엄격도, PM 요구사항 조율의 조합이 중요하다.",
        "멘토링 정책은 junior 성장과 senior 부하를 동시에 관리해야 효과가 지속된다.",
        "PM 역량은 allocation, bottleneck, requirement, scope control을 균형 있게 갖출 때 가장 안정적이다.",
    ])


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    section_overview(doc)
    section_model_system(doc)
    section_agents_tasks(doc)
    section_experiment_design(doc)
    section_results(doc)
    section_additional_results(doc)
    section_implications(doc)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
