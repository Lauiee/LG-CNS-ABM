#!/usr/bin/env python3
"""Build Korean DOCX report for calibrated F/G/H/I ABM experiment results."""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
REPORT_ROOT = Path(sys.argv[1]) if len(sys.argv) > 1 else sorted((BASE_DIR / "outputs" / "reports").glob("fghi_calibrated_report_*"))[-1]
NOTE_PATH = REPORT_ROOT / "report_summary.json"
DOCX_PATH = REPORT_ROOT / "LG_CNS_ABM_FGHI_실험결과_보고서.docx"

FONT_NAME = "Noto Sans CJK KR"
HEADING_BLUE = RGBColor(46, 116, 181)
HEADING_DARK = RGBColor(31, 77, 120)
INK = RGBColor(31, 36, 48)
MUTED = RGBColor(95, 102, 118)


def set_style_font(style, font_name=FONT_NAME, size=None, color=None, bold=None):
    font = style.font
    font.name = font_name
    if size is not None:
        font.size = Pt(size)
    if color is not None:
        font.color.rgb = color
    if bold is not None:
        font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, font_name=FONT_NAME, size=None, color=None, bold=None):
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=color, bold=bold)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=10.5, color=INK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    set_style_font(styles["Heading 1"], size=16, color=HEADING_BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)

    set_style_font(styles["Heading 2"], size=13, color=HEADING_BLUE, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)

    set_style_font(styles["Heading 3"], size=12, color=HEADING_DARK, bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    for style_name in ["List Bullet", "List Number"]:
        set_style_font(styles[style_name], size=10.5, color=INK)
        styles[style_name].paragraph_format.space_after = Pt(4)
        styles[style_name].paragraph_format.line_spacing = 1.167


def add_title(doc: Document, notes: dict):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("LG CNS ABM 시나리오 F/G/H/I 실험 결과 보고서")
    set_run_font(r, size=22, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("보정 후 100회 반복 실험 결과 기반 | runs=100, sprints=6")
    set_run_font(r, size=11, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    r = meta.add_run(f"생성 시각: {notes['meta']['generated_at']} | 데이터: calibrated validation summary CSV")
    set_run_font(r, size=9.5, color=MUTED)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=8.5, color=MUTED)


def add_image(doc: Document, image_path: str, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(6.35))
    add_caption(doc, caption)


def add_scope_table(doc: Document):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [1.0, 1.85, 1.85, 1.8]
    headers = ["시나리오", "실험 축", "주요 질문", "핵심 지표"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)
    rows = [
        ["F", "Project Type × Team Composition", "팀 구성별 성과와 비용 차이", "PR, 비용 효율, 재작업, 백로그"],
        ["G", "Project Type × PM Profile", "PM profile별 조정 효과", "CFR, rework, scope control, energy"],
        ["H", "Team Composition × Mentoring Intensity × Project Type", "멘토링 강도와 시니어 부하", "help resolution, bottleneck, rework"],
        ["I", "Requirement Volatility × PM Profile", "변동성 대응과 요구사항 조율", "rework, scope changes, Lead Time, CFR"],
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].width = Inches(widths[idx])
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_scenario_section(doc: Document, scenario_id: str, scenario: dict, figure_items: list[tuple[str, str]]):
    doc.add_heading(scenario["title"], level=1)

    doc.add_heading("실험 목적", level=2)
    add_bullets(doc, scenario["purpose"])

    doc.add_heading("의의", level=2)
    add_bullets(doc, scenario["meaning"])

    doc.add_heading("시각화 결과", level=2)
    for idx, (path, caption) in enumerate(figure_items, start=1):
        add_image(doc, path, f"그림 {scenario_id}-{idx}. {caption}")

    doc.add_heading("결과", level=2)
    add_bullets(doc, scenario["results"])

    doc.add_heading("해석", level=2)
    add_bullets(doc, scenario["interpretation"])


def figure_list(notes: dict, scenario_id: str) -> list[tuple[str, str]]:
    figures = notes["figures"][scenario_id]
    if scenario_id == "F":
        return [(figures["team_composition"]["png"], "프로젝트 유형과 팀 구성별 처리량, 비용 효율, 품질 부담, 시니어 부하 비교")]
    if scenario_id == "G":
        return [(figures["pm_profile"]["png"], "프로젝트 유형과 PM profile별 품질, 범위관리, 에너지, PM capacity 비교")]
    if scenario_id == "H":
        return [
            (figures["help_load"]["png"], "멘토링 강도별 도움 해결률과 시니어 멘토링 부하"),
            (figures["quality_energy"]["png"], "멘토링 강도별 시니어 병목, 재작업률, 평균 에너지"),
        ]
    if scenario_id == "I":
        return [
            (figures["rework_scope"]["png"], "요구사항 변동성별 재작업률, scope changes, scope changes prevented"),
            (figures["quality_energy"]["png"], "요구사항 변동성별 lead time, CFR, 평균 에너지"),
        ]
    return []


def add_appendix(doc: Document, notes: dict):
    doc.add_heading("보고서 읽는 법과 한계", level=1)
    add_bullets(doc, [
        "모든 수치는 calibrated validation run의 summary CSV에 있는 100회 반복 평균을 기준으로 작성했다.",
        "본 결과는 Mesa 기반 ABM 시뮬레이션의 구조와 파라미터 보정값에 따른 산출물이다.",
        "수치는 실제 조직 성과의 관측값이 아니라 모델 조건별 상대 비교 결과로 해석해야 한다.",
        "CFR은 Change Failure Rate이며, 값이 낮을수록 배포 품질 부담이 낮은 것으로 해석했다.",
        "rework rate, remaining backlog, senior mentoring load는 값이 높을수록 운영 부담이 큰 지표로 해석했다.",
        "PRs per Engineer, PRs per Cost, completed tasks per cost, avg_energy는 값이 높을수록 긍정적인 방향으로 해석했다.",
    ])

    doc.add_heading("산출물", level=1)
    add_bullets(doc, [
        f"원천 결과 폴더: {notes['meta']['source_root']}",
        f"보고서 산출 폴더: {notes['meta']['report_root']}",
        "시각화 파일: figures 폴더의 PNG 및 SVG 파일",
        "보고서 본문 데이터: report_summary.json",
    ])


def build_docx():
    notes = json.loads(NOTE_PATH.read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "LG CNS ABM 시나리오 F/G/H/I 실험 결과 보고서"
    doc.core_properties.subject = "Mesa 기반 ABM calibrated validation 결과 해석"
    doc.core_properties.author = "Codex"

    add_title(doc, notes)

    doc.add_heading("기술 요약", level=1)
    add_bullets(doc, notes["summary"])

    doc.add_page_break()
    doc.add_heading("실험 범위", level=1)
    add_scope_table(doc)

    for scenario_id in ["F", "G", "H", "I"]:
        add_scenario_section(
            doc,
            scenario_id,
            notes["scenarios"][scenario_id],
            figure_list(notes, scenario_id),
        )

    add_appendix(doc, notes)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
